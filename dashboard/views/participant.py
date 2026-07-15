from django.http import JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from dashboard.models import Participant, Training, Batch
from dashboard.forms import ParticipantForm
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from dashboard.models import Participant, Training, Batch
from docx import Document
from docx.shared import Inches
from django.db.models import Q
from django.contrib import messages
from django.contrib.auth.decorators import login_required

@login_required
def list_participants(request, training_id, batch_id):
    # training = get_object_or_404(Training, id=training_id)
    training = get_object_or_404(
        Training,
        id=training_id,
        created_by=request.user
        
    )
    batch = get_object_or_404(Batch, id=batch_id)
    participants = Participant.objects.filter(training=training, batch=batch, training__created_by=request.user)

    form = ParticipantForm()

    return render(request, "dashboard/list_participants.html", {
        "training": training,
        "batch": batch,
        "participants": participants,
        "form": form,
    })


#### commented on Aug
# def add_participant(request, training_id, batch_id, batch_number):
#     training = get_object_or_404(Training, pk=training_id)
#     batch = get_object_or_404(Batch, pk=batch_id)
#
#     if request.method == 'POST':
#         form = ParticipantForm(request.POST)
#         if form.is_valid():
#             official_id = form.cleaned_data.get("Official_ID")
#             exists = Participant.objects.filter(training=training, Official_ID=official_id).exists()
#
#             if exists:
#                 existing = Participant.objects.filter(training=training, Official_ID=official_id).first()
#                 messages.warning(request, f"This person is already enrolled in Batch #{existing.batch_number}")
#             else:
#                 participant = form.save(commit=False)
#                 participant.training = training
#                 participant.batch = batch
#                 participant.batch_number = batch_number
#                 participant.total_training_hours = batch.total_training_hours
#                 participant.save()
#                 messages.success(request, "✅ Participant added successfully.")
#                 return redirect('dashboard:participant_list', training_id=training.id, batch_id=batch.id)
#     else:
#         form = ParticipantForm()
#
#     return render(request, 'dashboard/add_participant.html', {
#         'form': form,
#         'training': training,
#         'batch': batch,
#         'batch_number': batch_number
#     })

def add_participant(request, training_id, batch_id, batch_number):
    # training = get_object_or_404(Training, pk=training_id)
    training = get_object_or_404(
        Training,
        id=training_id,
        created_by=request.user
    )
    batch = get_object_or_404(Batch, pk=batch_id)

    if request.method == 'POST':
        form = ParticipantForm(request.POST)
        if form.is_valid():
            # official_id = form.cleaned_data.get("Official_ID")

            # # Check for overlapping date conflicts
            # conflicts = Participant.objects.filter(
            #     Official_ID=official_id
            # ).filter(
            #     Q(batch__start_date__lte=batch.end_date) & Q(batch__end_date__gte=batch.start_date)
            # )

            # if conflicts.exists():
            #     conflict = conflicts.first()
            #     messages.error(
            #         request,
            #         f"This person is already attached with '{conflict.training.title}' training "
            #         f"from {conflict.batch.start_date} to {conflict.batch.end_date}."
            #     )
            #     return render(request, 'dashboard/add_participant.html', {
            #         'form': form,
            #         'training': training,
            #         'batch': batch,
            #         'batch_number': batch_number
            #     })

            # If no conflict, save participant
            participant = form.save(commit=False)
            participant.training = training
            participant.batch = batch
            participant.batch_number = batch_number
            participant.total_training_hours = batch.total_training_hours
            participant.save()
            messages.success(request, "✅ Participant added successfully.")
            return redirect('dashboard:participant_list', training_id=training.id, batch_id=batch.id)
    else:
        form = ParticipantForm()

    return render(request, 'dashboard/add_participant.html', {
        'form': form,
        'training': training,
        'batch': batch,
        'batch_number': batch_number
    })


def search_participant(request):
    official_id = request.GET.get('Official_ID', '')
    participant = Participant.objects.filter(Official_ID=official_id).first()

    if participant:
        return JsonResponse({
            'exists': True,
            'name': participant.name,
            'designation': participant.designation,
            'office_address': participant.office_address,
            'gender': participant.gender,
            'contact': participant.contact,
            'email': participant.email,
            'Official_ID': participant.Official_ID,
        })
    else:
        return JsonResponse({'exists': False})

# Update View
def update_participant(request, participant_id):
    participant = get_object_or_404(Participant, pk=participant_id, training__created_by=request.user)
    training = participant.training
    batch = participant.batch

    if request.method == 'POST':
        form = ParticipantForm(request.POST, instance=participant)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Participant updated successfully.")
            return redirect('dashboard:participant_list', training_id=training.id, batch_id=batch.id)
        else:
            messages.error(request, "❌ Please correct the errors.")
    else:
        form = ParticipantForm(instance=participant)

    return render(request, 'dashboard/update_participant.html', {
        'form': form,
        'participant': participant,
        'training': training,
        'batch': batch,
    })

# Delete View
def delete_participant(request, participant_id):
    participant = get_object_or_404(Participant, pk=participant_id, training__created_by=request.user)
    training_id = participant.training.id
    batch_id = participant.batch.id
    participant.delete()
    messages.success(request, "🗑️ Participant deleted successfully.")
    return redirect('dashboard:participant_list', training_id=training_id, batch_id=batch_id)

# Utility function to convert English digits to Bangla
def convert_to_bangla_number(number):
    en = "0123456789"
    bn = "০১২৩৪৫৬৭৮৯"
    trans_table = str.maketrans(en, bn)
    return str(number).translate(trans_table)

def generate_participant_word(request, training_id, batch_id):
    # training = get_object_or_404(Training, pk=training_id)
    training = get_object_or_404(
        Training,
        id=training_id,
        created_by=request.user
    )
    batch = get_object_or_404(Batch, pk=batch_id)
    participants = Participant.objects.filter(training_id=training_id, batch_id=batch_id, training__created_by=request.user)

    doc = Document()
    doc.add_heading('প্রশিক্ষণার্থীদের তালিকা', level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ক্রমিক'
    hdr_cells[1].text = 'নাম'
    hdr_cells[2].text = 'পদবি'
    hdr_cells[3].text = 'কার্যালয়'
    hdr_cells[4].text = 'মোবাইল'

    for i, p in enumerate(participants, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = convert_to_bangla_number(i)
        row_cells[1].text = p.name
        row_cells[2].text = p.designation
        row_cells[3].text = p.office_address
        row_cells[4].text = p.contact or ""

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Participants_List_Training_{training_id}_Batch_{batch_id}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    doc.save(response)
    return response
