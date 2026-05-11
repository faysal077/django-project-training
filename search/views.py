from django.shortcuts import render
from django.shortcuts import render
from dashboard.models import Training, Participant
from django.http import JsonResponse
from django.template.loader import render_to_string
from django.shortcuts import render
from django.template.loader import render_to_string
from django.http import JsonResponse, HttpResponse
from django.db.models import Q
from django.contrib.auth.decorators import login_required

@login_required
def search_dashboard(request):
    #return render(request, 'search/_search_dashboard_template.html')
    return render(request, 'basee.html')


from django.shortcuts import render
from django.http import JsonResponse
from dashboard.models import Participant


@login_required
def search_by_name(request):
    participants = []
    participant_name = ''
    searched = False

    if request.method == 'POST':
        participant_name = request.POST.get('participant_name', '').strip()
        searched = True

        if participant_name:
            results = Participant.objects.filter(
                name__icontains=participant_name
            ).select_related('training', 'batch') \
             .values(
                'name',
                'Official_ID',
                'designation',
                'office_address',
                'training__title',
                'batch__start_date',
                'batch__end_date',
                'batch__batch_number'
             ).order_by('-batch__start_date')

            # Normalize keys for template
            participants = [
                {
                    'name': r['name'],
                    'Official_ID': r['Official_ID'],
                    'designation': r['designation'],
                    'office_address': r['office_address'],
                    'training_title': r['training__title'],
                    'start_date': r['batch__start_date'],
                    'end_date': r['batch__end_date'],
                    'batch_number': r['batch__batch_number'],
                }
                for r in results
            ]

    return render(request, 'search/search_by_name_template.html', {
        'participants': participants,
        'participant_name': participant_name,
        'searched': searched
    })



'''
def search_by_name(request):
    trainings = []
    name = ''
    searched = False

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        searched = True
        if name:
            trainings = Participant.objects.filter(name__icontains=name).select_related('training', 'batch') \
                .values(
                'training__title',
                'batch__start_date',
                'batch__end_date',
                'batch__batch_number'
            ).order_by('-batch__start_date')

            # Rename field keys to match template expectations
            trainings = [
                {
                    'title': t['training__title'],
                    'start_date': t['batch__start_date'],
                    'end_date': t['batch__end_date'],
                    'batch_number': t['batch__batch_number']
                }
                for t in trainings
            ]

    return render(request, 'search/search_by_name_template.html', {
        'trainings': trainings,
        'name': name,
        'searched': searched,
    })

'''
@login_required
def search_by_training(request):
    participants = []
    training_name = ''
    searched = False

    if request.method == 'POST':
        training_name = request.POST.get('training_name', '').strip()
        searched = True

        if training_name:
            results = Participant.objects.filter(
                training__title__icontains=training_name
            ).select_related('training', 'batch') \
                .values(
                'name',
                'designation',
                'batch__start_date',
                'batch__end_date',
                'batch__batch_number'
            ).order_by('-batch__start_date')

            # Normalize keys for template
            participants = [
                {
                    'name': r['name'],
                    'designation': r['designation'],
                    'start_date': r['batch__start_date'],
                    'end_date': r['batch__end_date'],
                    'batch_number': r['batch__batch_number'],
                } for r in results
            ]

    return render(request, 'search/search_by_training_template.html', {
        'participants': participants,
        'training_name': training_name,
        'searched': searched
    })
@login_required
def search_not_taken(request):
    participants = []
    training_name = ''
    searched = False

    if request.method == 'POST':
        training_name = request.POST.get('training_name', '').strip()
        searched = True

        if training_name:
            # Find the training object (assuming title is unique or using icontains)
            training_qs = Training.objects.filter(title__icontains=training_name)

            if training_qs.exists():
                training = training_qs.first()

                # Get IDs of participants who attended the training
                attended_ids = Participant.objects.filter(training=training).values_list('id', flat=True)

                # Get all participants who did NOT attend this training
                participants = Participant.objects.exclude(id__in=attended_ids)

    return render(request, 'search/search_not_taken_template.html', {
        'participants': participants,
        'training_name': training_name,
        'searched': searched
    })
'''
def search_not_taken(request):
    trainings = []
    name = ''
    searched = False

    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        searched = True

        if name:
            # Find all training IDs attended by the participant
            attended_ids = Participant.objects.filter(name__icontains=name).values_list('training_id', flat=True)

            # Get trainings not attended
            trainings = Training.objects.exclude(id__in=attended_ids).order_by('-created_at') \
                .values('title', 'created_at', 'organizer', 'training_type', 'id')  # adjust if needed

            # Optionally: convert created_at to start_date if you have a batch model for that info

    return render(request, 'search/search_not_taken_template.html', {
        'name': name,
        'searched': searched,
        'trainings': trainings,
    })

'''
from django.shortcuts import render
from dashboard.models import Participant

@login_required
def search_by_multiple_trainings(request):
    participants = []
    training_names = ''
    searched = False

    if request.method == 'POST':
        training_names = request.POST.get('training_names', '')
        searched = True
        training_list = [t.strip() for t in training_names.split(',') if t.strip()]

        if training_list:
            results = Participant.objects.filter(
                training__title__in=training_list
            ).select_related('training') \
                .values(
                'name',
                'designation',
                'batch_number',
                'training__title'
            )

            # Normalize keys for template
            participants = [
                {
                    'name': r['name'],
                    'designation': r['designation'],
                    'batch_number': r['batch_number'],
                    'training_name': r['training__title'],
                } for r in results
            ]

    return render(request, 'search/search_by_multiple_template.html', {
        'training_names': training_names,
        'participants': participants,
        'searched': searched
    })

from datetime import date
from django.db.models import Min, Max
from django.contrib.auth.decorators import login_required
from dashboard.models import Participant, Batch


@login_required
def search_by_timeline(request):

    fiscal_year = request.GET.get("fiscal_year")
    month = request.GET.get("month")
    training_type = request.GET.get("training_type")

    # ✅ Start with base queryset ONCE
    participants = Participant.objects.select_related("training", "batch")

    # -------------------------
    # Fiscal Year Filter
    # -------------------------
    if fiscal_year:
        start_year, end_year = map(int, fiscal_year.split("-"))

        fiscal_start = date(start_year, 7, 1)
        fiscal_end = date(end_year, 6, 30)

        participants = participants.filter(
            batch__start_date__range=(fiscal_start, fiscal_end)
        )

    else:
        # ❗ If no fiscal year selected → return empty result
        participants = Participant.objects.none()

    # -------------------------
    # Training Type Filter
    # -------------------------
    if training_type and training_type != "all":
        participants = participants.filter(
            training__training_type=training_type
        )

    # -------------------------
    # Month Filter
    # -------------------------
    if month and month != "all":
        participants = participants.filter(
            batch__start_date__month=int(month)
        )

    # -------------------------
    # Sorting
    # -------------------------
    participants = participants.order_by(
        "-batch__start_date",
        "training__title"
    )

    # -------------------------
    # Totals
    # -------------------------
    total_trainings = participants.values("training").distinct().count()
    total_batches = participants.values("batch").distinct().count()
    total_participants = participants.count()

    # -------------------------
    # Fiscal year list (unchanged)
    # -------------------------
    batch_dates = Batch.objects.aggregate(
        min_date=Min("start_date"),
        max_date=Max("start_date")
    )

    fiscal_years = []

    if batch_dates["min_date"] and batch_dates["max_date"]:
        min_year = batch_dates["min_date"].year - 1
        max_year = batch_dates["max_date"].year + 1

        for y in range(min_year, max_year):
            fiscal_years.append(f"{y}-{y+1}")

    context = {
        "participants": participants,
        "total_trainings": total_trainings,
        "total_batches": total_batches,
        "total_participants": total_participants,
        "selected_fiscal_year": fiscal_year,
        "selected_month": month if month else "all",
        "selected_training_type": training_type if training_type else "all",
        "training_types": [
            ('ইন-হাউজ/অভ্যন্তরীণ', 'ইন-হাউজ/অভ্যন্তরীণ'),
            ('স্থানীয়', 'স্থানীয়'),
            ('বৈদেশিক', 'বৈদেশিক'),
        ],
        "fiscal_years": fiscal_years,
    }

    return render(request, "search/search_by_timeline.html", context)

from docx import Document
from django.http import HttpResponse
from datetime import date


@login_required
def download_word(request):

    fiscal_year = request.GET.get("fiscal_year")
    month = request.GET.get("month")
    training_type = request.GET.get("training_type")

    participants = Participant.objects.select_related("training", "batch")

    # -------------------------
    # Apply Filters (same as UI)
    # -------------------------
    if fiscal_year:
        start_year, end_year = map(int, fiscal_year.split("-"))

        fiscal_start = date(start_year, 7, 1)
        fiscal_end = date(end_year, 6, 30)

        participants = participants.filter(
            batch__start_date__range=(fiscal_start, fiscal_end)
        )

    if month and month != "all":
        participants = participants.filter(
            batch__start_date__month=int(month)
        )

    if training_type and training_type != "all":
        participants = participants.filter(
            training__training_type=training_type
        )

    participants = participants.order_by(
        "-batch__start_date",
        "training__title"
    )

    # -------------------------
    # Create Word Document
    # -------------------------
    doc = Document()
    doc.add_heading('প্রশিক্ষণার্থীদের তালিকা', 0)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ক্রমিক'
    hdr_cells[1].text = 'প্রশিক্ষণ'
    hdr_cells[2].text = 'ব্যাচ'
    hdr_cells[3].text = 'নাম'
    hdr_cells[4].text = 'পদবি'
    hdr_cells[5].text = 'কার্যালয়'
    hdr_cells[6].text = 'প্রশিক্ষণ শুরুর তারিখ'

    # -------------------------
    # Fill Data
    # -------------------------
    for i, p in enumerate(participants, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)  # or convert_to_bangla_number(i)
        row_cells[1].text = p.training.title
        row_cells[2].text = f"Batch {p.batch.batch_number}"
        row_cells[3].text = p.name
        row_cells[4].text = p.designation
        row_cells[5].text = p.office_address
        row_cells[6].text = p.batch.start_date.strftime("%d-%m-%Y")

    # -------------------------
    # Response
    # -------------------------
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    filename = "Training_Report.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    doc.save(response)
    return response
